import tempfile
import unittest
from pathlib import Path

from docx import Document

from osah.application.services.create_port_site_passport import create_port_site_passport
from osah.application.services.export_port_shift_briefing_to_docx import export_port_shift_briefing_to_docx
from osah.application.services.initialize_application import initialize_application
from osah.domain.entities.access_role import AccessRole
from osah.domain.entities.port_site_passport_input import PortSitePassportInput
from osah.infrastructure.config.application_paths import build_application_paths
from osah.infrastructure.logging.shutdown_logging import shut_down_logging


class ExportPortShiftBriefingToDocxTests(unittest.TestCase):
    """Тести експорту оперативного листа зміни ПОРТ-Р у .docx.
    Tests for exporting the PORT-R shift briefing to a .docx file.
    """

    def test_export_creates_docx_with_expected_sections(self) -> None:
        """Створює реальний docx-файл і всі очікувані секції присутні в тексті.
        Creates a real docx file and all expected sections are present in the text.
        """

        with tempfile.TemporaryDirectory() as temporary_directory:
            project_root = Path(temporary_directory)
            context = initialize_application(build_application_paths(project_root))
            try:
                passport_input = _make_passport_input("P-EXP-1")
                passport_id = create_port_site_passport(
                    context.database_path,
                    passport_input,
                    access_role=AccessRole.INSPECTOR,
                )

                export_result = export_port_shift_briefing_to_docx(
                    context.database_path,
                    project_root,
                    passport_id,
                    actor_name="tester",
                    access_role=AccessRole.INSPECTOR,
                )

                self.assertTrue(export_result.file_path.exists())
                self.assertGreater(export_result.file_path.stat().st_size, 0)
                self.assertEqual(
                    export_result.file_path.parent,
                    project_root / "exports" / "port_r",
                )

                document = Document(str(export_result.file_path))
                full_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
                expected_section_titles = (
                    "Паспорт ділянки P-EXP-1",
                    "1. ЗАГАЛЬНІ ДАНІ",
                    "2. БАЗОВИЙ ПРОФІЛЬ ДІЛЯНКИ",
                    "3. КЛЮЧОВІ РИЗИКИ ДІЛЯНКИ",
                    "4. ПЕРЕВІРКА Т-П-С-В-Б ПЕРЕД ПОЧАТКОМ РОБІТ",
                    "5. ПЕРЕВІРКА КРИТИЧНИХ БАР'ЄРІВ",
                    "6. ФАКТИЧНІ ЗМІНИ ПІД ЧАС ЗМІНИ",
                    "7. ПІДСУМОК ЗМІНИ",
                    "8. ПІДПИС ВІДПОВІДАЛЬНОГО ЗМІНИ",
                )
                for title in expected_section_titles:
                    self.assertIn(title, full_text)

                tables_text = "\n".join(
                    cell.text for table in document.tables for row in table.rows for cell in row.cells
                )
                self.assertIn("P-EXP-1", full_text)
                self.assertIn("Причал №7", tables_text)
            finally:
                shut_down_logging()


def _make_passport_input(passport_code: str) -> PortSitePassportInput:
    return PortSitePassportInput(
        passport_code=passport_code,
        site_name="Причал №7",
        site_type="ВРР",
        site_location="Одеський порт",
        site_description="",
        work_kind="вантажні роботи",
        typical_operations="стропування, переміщення",
        work_mode="одна зміна",
        typical_cargo="метал",
        cargo_features="",
        main_equipment="портовий кран",
        lifting_devices="стропи, траверса",
        has_railway_zone=False,
        has_auto_zone=True,
        has_crane_zone=True,
        crew_composition="бригада 6 осіб",
        responsible_person="майстер зміни",
        has_contractors=False,
        contractors_note="",
        zone_kind="відкритий причал",
        has_night_works=False,
        weather_features="",
        has_limited_visibility=False,
        has_height_work=False,
        has_water_edge_work=True,
        has_stack_edge_work=False,
        has_communication_barrier=True,
        communication_barrier="радіостанції",
        has_fencing_barrier=True,
        fencing_barrier="конуси",
        has_signalman=True,
        has_lighting_barrier=True,
        lighting_barrier="прожектори",
        ppe_text="каска, жилет, рукавиці",
        additional_barriers="",
    )


if __name__ == "__main__":
    unittest.main()
